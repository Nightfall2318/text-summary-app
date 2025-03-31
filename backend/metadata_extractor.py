"""
File Metadata Extraction Module
This module provides functions to extract metadata from various file types.
"""

import os
import datetime
from PIL import Image
from PIL.ExifTags import TAGS
from pypdf import PdfReader
import docx
import magic
import mimetypes
import zipfile
import xml.etree.ElementTree as ET
import dateutil.parser

def get_basic_metadata(file_path):
    """Extract basic file metadata available for all file types"""
    try:
        # Get file stats
        stats = os.stat(file_path)
        
        # Use libmagic to determine MIME type
        mime = magic.Magic(mime=True)
        mime_type = mime.from_file(file_path)
        
        # Get basic metadata
        file_name = os.path.basename(file_path)
        file_extension = os.path.splitext(file_path)[1].lower()
        
        # Basic metadata dictionary
        metadata = {
            "filename": file_name,
            "extension": file_extension,
            "mime_type": mime_type,
            "size_bytes": stats.st_size,
            "size_formatted": format_file_size(stats.st_size),
            "created_time": format_timestamp(stats.st_ctime),
            "modified_time": format_timestamp(stats.st_mtime),
            "accessed_time": format_timestamp(stats.st_atime),
            "file_path": file_path
        }
        
        return metadata
    except Exception as e:
        print(f"Error extracting basic metadata: {e}")
        return {"error": str(e)}

def extract_metadata(file_path):
    """Extract metadata based on file type"""
    # Get basic metadata first
    metadata = get_basic_metadata(file_path)
    
    # Extract specific metadata based on file extension
    file_extension = os.path.splitext(file_path)[1].lower()
    
    try:
        if file_extension == '.pdf':
            pdf_metadata = extract_pdf_metadata(file_path)
            metadata.update(pdf_metadata)
            
        elif file_extension == '.docx':
            docx_metadata = extract_docx_metadata(file_path)
            metadata.update(docx_metadata)
            
        elif file_extension in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp', '.gif']:
            image_metadata = extract_image_metadata(file_path)
            metadata.update(image_metadata)
            
        elif file_extension == '.txt':
            txt_metadata = extract_text_metadata(file_path)
            metadata.update(txt_metadata)
            
        # Add more file types as needed
    except Exception as e:
        metadata["metadata_error"] = str(e)
    
    return metadata

def extract_pdf_metadata(file_path):
    """Extract metadata from PDF files"""
    try:
        with open(file_path, 'rb') as file:
            reader = PdfReader(file)
            
            # Initialize metadata dict
            metadata = {
                "type": "PDF Document",
                "pages": len(reader.pages),
                "encrypted": reader.is_encrypted
            }
            
            # Extract document info if available
            if reader.metadata:
                info = reader.metadata
                
                # Map common metadata fields
                field_mapping = {
                    "/Title": "title",
                    "/Author": "author",
                    "/Subject": "subject",
                    "/Keywords": "keywords",
                    "/Creator": "creator",
                    "/Producer": "producer",
                    "/CreationDate": "creation_date",
                    "/ModDate": "modification_date"
                }
                
                # Process each field
                for field, mapped_name in field_mapping.items():
                    value = info.get(field)
                    if value:
                        # Handle date fields specially
                        if field in ["/CreationDate", "/ModDate"] and isinstance(value, str):
                            # PDF dates are in format: D:YYYYMMDDHHmmSSOHH'mm'
                            # Try to parse it if possible
                            try:
                                if value.startswith("D:"):
                                    value = value[2:]  # Remove D: prefix
                                    # Basic parsing for common format
                                    year = value[0:4]
                                    month = value[4:6]
                                    day = value[6:8]
                                    parsed_date = f"{year}-{month}-{day}"
                                    if len(value) > 8:
                                        hour = value[8:10]
                                        minute = value[10:12]
                                        parsed_date += f" {hour}:{minute}"
                                    metadata[mapped_name] = parsed_date
                                else:
                                    metadata[mapped_name] = value
                            except:
                                metadata[mapped_name] = value
                        else:
                            metadata[mapped_name] = value
            
            # Estimate word count from first page for a rough idea
            if len(reader.pages) > 0:
                first_page_text = reader.pages[0].extract_text()
                words_per_page = len(first_page_text.split())
                estimated_words = words_per_page * len(reader.pages)
                metadata["estimated_word_count"] = estimated_words
            
            return metadata
    except Exception as e:
        print(f"Error extracting PDF metadata: {e}")
        return {"pdf_metadata_error": str(e)}

def extract_docx_metadata(file_path):
    """Extract metadata from DOCX files"""
    try:
        doc = docx.Document(file_path)
        
        # Basic document properties
        metadata = {
            "type": "Word Document",
            "paragraph_count": len(doc.paragraphs),
            "word_count": count_words_in_docx(doc),
        }
        
        # Try to extract core properties from the Office Open XML format
        try:
            with zipfile.ZipFile(file_path) as docx_zip:
                if 'docProps/core.xml' in docx_zip.namelist():
                    with docx_zip.open('docProps/core.xml') as core_xml:
                        tree = ET.parse(core_xml)
                        root = tree.getroot()
                        
                        # Define XML namespaces
                        namespaces = {
                            'dc': 'http://purl.org/dc/elements/1.1/',
                            'dcterms': 'http://purl.org/dc/terms/',
                            'cp': 'http://schemas.openxmlformats.org/package/2006/metadata/core-properties'
                        }
                        
                        # Extract common properties
                        title = root.find('.//dc:title', namespaces)
                        if title is not None and title.text:
                            metadata['title'] = title.text
                            
                        creator = root.find('.//dc:creator', namespaces)
                        if creator is not None and creator.text:
                            metadata['author'] = creator.text
                            
                        subject = root.find('.//dc:subject', namespaces)
                        if subject is not None and subject.text:
                            metadata['subject'] = subject.text
                            
                        description = root.find('.//dc:description', namespaces)
                        if description is not None and description.text:
                            metadata['description'] = description.text
                            
                        created = root.find('.//dcterms:created', namespaces)
                        if created is not None and created.text:
                            try:
                                date = dateutil.parser.parse(created.text)
                                metadata['creation_date'] = date.strftime('%Y-%m-%d %H:%M:%S')
                            except:
                                metadata['creation_date'] = created.text
                                
                        modified = root.find('.//dcterms:modified', namespaces)
                        if modified is not None and modified.text:
                            try:
                                date = dateutil.parser.parse(modified.text)
                                metadata['modification_date'] = date.strftime('%Y-%m-%d %H:%M:%S')
                            except:
                                metadata['modification_date'] = modified.text
                
                # Also try to get the app properties
                if 'docProps/app.xml' in docx_zip.namelist():
                    with docx_zip.open('docProps/app.xml') as app_xml:
                        tree = ET.parse(app_xml)
                        root = tree.getroot()
                        
                        # App properties namespace
                        ns = {'ep': 'http://schemas.openxmlformats.org/officeDocument/2006/extended-properties'}
                        
                        # Extract app properties
                        company = root.find('.//ep:Company', ns)
                        if company is not None and company.text:
                            metadata['company'] = company.text
                            
                        app_version = root.find('.//ep:AppVersion', ns)
                        if app_version is not None and app_version.text:
                            metadata['app_version'] = app_version.text
                            
                        # More accurate word count
                        word_count = root.find('.//ep:Words', ns)
                        if word_count is not None and word_count.text:
                            try:
                                metadata['word_count'] = int(word_count.text)
                            except:
                                # Keep the estimated count if conversion fails
                                pass
        except Exception as e:
            metadata["docx_xml_metadata_error"] = str(e)
        
        return metadata
    except Exception as e:
        print(f"Error extracting DOCX metadata: {e}")
        return {"docx_metadata_error": str(e)}

def extract_image_metadata(file_path):
    """Extract metadata from image files"""
    try:
        with Image.open(file_path) as img:
            # Basic image properties
            metadata = {
                "type": "Image",
                "format": img.format,
                "mode": img.mode,
                "width": img.width,
                "height": img.height,
                "resolution": img.info.get('dpi', 'Unknown')
            }
            
            # Extract EXIF data if available
            if hasattr(img, '_getexif') and img._getexif():
                exif = img._getexif()
                if exif:
                    exif_data = {}
                    for tag_id, value in exif.items():
                        tag = TAGS.get(tag_id, tag_id)
                        # Handle special cases like GPS data
                        if tag == 'GPSInfo':
                            # Simplified GPS handling
                            exif_data['GPS Data Present'] = 'Yes'
                        else:
                            # Skip binary data and convert bytes to strings
                            if isinstance(value, bytes):
                                try:
                                    value = value.decode('utf-8')
                                except:
                                    value = str(value)
                            
                            # Add tag and value to the dictionary
                            exif_data[tag] = value
                    
                    # Store common EXIF data if available
                    common_exif = [
                        ('Make', 'camera_make'),
                        ('Model', 'camera_model'),
                        ('DateTime', 'date_time'),
                        ('ExposureTime', 'exposure_time'),
                        ('FNumber', 'f_number'),
                        ('ISOSpeedRatings', 'iso_speed'),
                        ('FocalLength', 'focal_length'),
                        ('Software', 'software')
                    ]
                    
                    for exif_tag, meta_key in common_exif:
                        if exif_tag in exif_data:
                            metadata[meta_key] = exif_data[exif_tag]
                    
                    # Store all EXIF data in a separate field
                    metadata['exif'] = exif_data
            
            return metadata
    except Exception as e:
        print(f"Error extracting image metadata: {e}")
        return {"image_metadata_error": str(e)}

def extract_text_metadata(file_path):
    """Extract metadata from text files"""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
            content = file.read()
            
            lines = content.split('\n')
            words = content.split()
            
            metadata = {
                "type": "Text Document",
                "line_count": len(lines),
                "word_count": len(words),
                "character_count": len(content),
                "character_count_no_spaces": len(content.replace(' ', ''))
            }
            
            # Try to detect encoding
            try:
                import chardet
                with open(file_path, 'rb') as raw_file:
                    result = chardet.detect(raw_file.read())
                    metadata['encoding'] = result['encoding']
                    metadata['encoding_confidence'] = f"{result['confidence']:.2%}"
            except ImportError:
                # chardet not available
                pass
                
            return metadata
    except UnicodeDecodeError:
        # Try again with latin-1 encoding for binary files
        try:
            with open(file_path, 'r', encoding='latin-1', errors='replace') as file:
                content = file.read()
                metadata = {
                    "type": "Text Document (Non-UTF8)",
                    "encoding": "unknown",
                    "line_count": len(content.split('\n')),
                    "word_count": len(content.split()),
                    "character_count": len(content)
                }
                return metadata
        except Exception as inner_e:
            return {"text_metadata_error": str(inner_e)}
    except Exception as e:
        print(f"Error extracting text metadata: {e}")
        return {"text_metadata_error": str(e)}

def count_words_in_docx(doc):
    """Count words in a docx document"""
    word_count = 0
    for para in doc.paragraphs:
        word_count += len(para.text.split())
    return word_count

def format_timestamp(timestamp):
    """Convert Unix timestamp to readable date time format"""
    return datetime.datetime.fromtimestamp(timestamp).strftime('%Y-%m-%d %H:%M:%S')

def format_file_size(size_bytes):
    """Format file size in a human-readable format"""
    if size_bytes < 1024:
        return f"{size_bytes} B"
    elif size_bytes < 1024 * 1024:
        return f"{size_bytes / 1024:.2f} KB"
    elif size_bytes < 1024 * 1024 * 1024:
        return f"{size_bytes / (1024 * 1024):.2f} MB"
    else:
        return f"{size_bytes / (1024 * 1024 * 1024):.2f} GB"